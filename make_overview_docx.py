"""
make_overview_docx.py -- Renders PROJECT_OVERVIEW.md as a .docx.

PROJECT_OVERVIEW.md is the single source of truth for the plain-English
overview (make_overview_pdf.py is a separately hand-authored PDF with its
own duplicated content). This script instead mechanically converts the
Markdown file itself, so the .docx can't drift from PROJECT_OVERVIEW.md the
way two independently hand-written documents can drift from each other.

Supports the Markdown actually used in that file: #/##/### headings, **bold**,
*italic*, `code` spans, "- " bullets (with wrapped continuation lines), pipe
tables, ```-fenced code blocks, and "---" horizontal rules. Not a general
Markdown parser -- just enough for this one file.

Run:
    python make_overview_docx.py
Output: documents/Project_Overview.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).parent
SRC_PATH = PROJECT_ROOT / "PROJECT_OVERVIEW.md"
OUT_PATH = PROJECT_ROOT / "documents" / "Project_Overview.docx"

NAVY = RGBColor(30, 58, 95)
BLUE = RGBColor(37, 99, 235)
GREY = RGBColor(90, 100, 115)
TEXT = RGBColor(30, 41, 59)
LIGHT = "EFF6FF"
CODEBG = "F4F6F8"
FONT = "Calibri"
MONO = "Consolas"

_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def shade(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)


def add_inline(paragraph, text, base_size=11, base_color=TEXT):
    """Split text on **bold** / *italic* / `code` and add runs accordingly."""
    for chunk in _TOKEN_RE.split(text):
        if not chunk:
            continue
        r = paragraph.add_run()
        r.font.name = FONT
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color
        if chunk.startswith("**") and chunk.endswith("**"):
            r.text = chunk[2:-2]
            r.bold = True
            r.font.color.rgb = NAVY
        elif chunk.startswith("`") and chunk.endswith("`"):
            r.text = chunk[1:-1]
            r.font.name = MONO
            r.font.color.rgb = RGBColor(180, 40, 40)
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            r.text = chunk[1:-1]
            r.italic = True
            r.font.color.rgb = BLUE
        else:
            r.text = chunk


def h_heading(doc, text, level):
    text = text.strip()
    text = re.sub(r"^[^\w(]+", "", text).strip()  # drop leading emoji
    p = doc.add_paragraph()
    if level == 2:
        shade(p, LIGHT)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = NAVY
        r.font.name = FONT
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = BLUE
        r.font.name = FONT


def add_paragraph(doc, text, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bullet:
        p.style = doc.styles["List Bullet"]
    add_inline(p, text)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    shade(p, CODEBG)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(9.5)
        r.font.color.rgb = TEXT
        if i < len(lines) - 1:
            p.add_run().add_break()


def split_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_table(doc, rows):
    header, *body_rows = rows
    n_cols = len(header)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        p = cell.paragraphs[0]
        shade(p, LIGHT)
        add_inline(p, text, base_size=10.5)
        for run in p.runs:
            run.bold = True
    for row_vals in body_rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_vals):
            p = cell.paragraphs[0]
            add_inline(p, text, base_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build():
    lines = SRC_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    para_buf = []
    bullet_buf = []

    def flush_para():
        if para_buf:
            add_paragraph(doc, " ".join(para_buf))
            para_buf.clear()

    def flush_bullet():
        if bullet_buf:
            add_paragraph(doc, " ".join(bullet_buf), bullet=True)
            bullet_buf.clear()

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped == "":
            flush_para()
            flush_bullet()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_para(); flush_bullet()
            title = re.sub(r"^#\s+", "", stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(24)
            r.font.color.rgb = NAVY
            r.font.name = FONT
            i += 1
            continue

        if stripped.startswith("## "):
            flush_para(); flush_bullet()
            h_heading(doc, stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_para(); flush_bullet()
            h_heading(doc, stripped[4:], level=3)
            i += 1
            continue

        if stripped == "---":
            flush_para(); flush_bullet()
            hr(doc)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_para(); flush_bullet()
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code_lines)
            continue

        if stripped.startswith("|"):
            flush_para(); flush_bullet()
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [split_table_row(l) for l in table_lines
                    if not re.match(r"^\|[\s\-|]+\|$", l)]
            add_table(doc, rows)
            continue

        if stripped.startswith("- "):
            flush_para()
            if bullet_buf:
                flush_bullet()
            bullet_buf.append(stripped[2:])
            i += 1
            continue

        if bullet_buf and (raw.startswith("  ") or raw.startswith("\t")):
            bullet_buf.append(stripped)
            i += 1
            continue

        flush_bullet()
        para_buf.append(stripped)
        i += 1

    flush_para()
    flush_bullet()

    OUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build()
