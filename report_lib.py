"""
report_lib.py — Formatting helpers for the full project report.

Implements the house style required by PROJECT TEMPLATE_Civil.docx:

  * A4 paper, 2 cm margins on all four sides
  * Times New Roman 12 pt throughout, double line spacing, justified body text
  * Chapter headings centred and bold; first-level headings in Title Case and
    bold; second-level headings in Sentence case and bold (the template forbids
    a third level, so the document never uses one)
  * Table titles above tables, figure captions below figures
  * APA 6 reference list with a hanging indent
  * Roman numerals for the front matter, Arabic numerals for the body

It also renders Python source as a shaded, syntax-coloured monospace block for
the appendix, and extracts that source from the real project modules so the
listing can never drift from the code that produced the results.
"""

import ast
import builtins
import io
import keyword
import tokenize
from pathlib import Path

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── house style constants ────────────────────────────────────────────────────
FONT = "Times New Roman"
BODY_PT = Pt(12)
CODE_FONT = "Consolas"
CODE_PT = Pt(8.5)  # ~99 characters fit the 16.4 cm code column without wrapping

A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)
MARGIN = Cm(2.0)

# Syntax colours for the appendix listings (light background, print-safe)
CLR_KEYWORD = RGBColor(0x00, 0x00, 0xC0)
CLR_STRING = RGBColor(0xA3, 0x15, 0x15)
CLR_COMMENT = RGBColor(0x1E, 0x7A, 0x1E)
CLR_NUMBER = RGBColor(0x09, 0x86, 0x58)
CLR_DEFNAME = RGBColor(0x79, 0x5E, 0x26)
CLR_CLASSNAME = RGBColor(0x26, 0x7F, 0x99)
CLR_BUILTIN = RGBColor(0x26, 0x7F, 0x99)
CLR_PLAIN = RGBColor(0x1F, 0x1F, 0x1F)
CODE_SHADE = "F4F4F4"

BUILTIN_NAMES = set(dir(builtins))


# ── low-level XML helpers ────────────────────────────────────────────────────

def _shade_paragraph(par, fill_hex):
    """Apply a solid background fill to a paragraph."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    par._p.get_or_add_pPr().append(shd)


def _field(par, instruction, cached_text=""):
    """Insert a Word field. `cached_text` is what shows before F9 refreshes it."""
    r1 = par.add_run()._r
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r1.append(fc)

    r2 = par.add_run()._r
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instruction
    r2.append(it)

    r3 = par.add_run()._r
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "separate")
    r3.append(fc)

    run = par.add_run(cached_text)
    run.font.name = FONT
    run.font.size = BODY_PT

    r5 = par.add_run()._r
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r5.append(fc)


def _page_number_format(section, fmt, start=None):
    """Set the page-number format ('lowerRoman' / 'decimal') for a section."""
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    section._sectPr.append(pg)


# ── document / section setup ─────────────────────────────────────────────────

def setup_page(section):
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN


def setup_styles(doc):
    """Base every built-in style the document uses on Times New Roman 12 pt."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_PT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for name, size, bold, italic in (
        ("Heading 1", 12, True, False),
        ("Heading 2", 12, True, False),
        ("Heading 3", 12, True, False),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = italic
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # Dedicated caption styles so the List of Figures and List of Tables can be
    # generated as two separate, refreshable TOC fields.
    for name in ("FigureCaption", "TableCaption"):
        try:
            st = doc.styles.add_style(name, 1)  # 1 = paragraph style
        except ValueError:
            st = doc.styles[name]
        st.base_style = doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(11)
        st.font.bold = False
        st.quick_style = False


def footer_page_number(section, linked=False):
    section.footer.is_linked_to_previous = linked
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    _field(par, " PAGE ", "")


def clear_footer(section):
    section.footer.is_linked_to_previous = False
    par = section.footer.paragraphs[0]
    for r in list(par.runs):
        r._r.getparent().remove(r._r)


def new_section(doc, page_fmt=None, start=None, numbered=True):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(sec)
    if page_fmt:
        _page_number_format(sec, page_fmt, start)
    if numbered:
        footer_page_number(sec)
    else:
        clear_footer(sec)
    return sec


# ── text blocks ──────────────────────────────────────────────────────────────

def _spacing(par, line=2.0, before=0, after=6):
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return par


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False, bold=False):
    """A double-spaced body paragraph."""
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = bold
    par.alignment = align
    _spacing(par)
    if indent:
        par.paragraph_format.first_line_indent = Cm(1.27)
    return par


def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(par, line=1.5, after=2)
    return par


def chapter_heading(doc, number_word, title, page_break=True):
    """'CHAPTER ONE' / 'INTRODUCTION', centred and bold, on a fresh page."""
    if page_break:
        doc.add_page_break()
    par = doc.add_paragraph(style="Heading 1")
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(f"CHAPTER {number_word}")
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    _spacing(par, line=2.0, before=0, after=0)

    par2 = doc.add_paragraph(style="Heading 1")
    par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par2.add_run(title.upper())
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    _spacing(par2, line=2.0, before=0, after=12)
    return par2


def h1(doc, text):
    """First-level heading, e.g. '2.1 Classification of Forecasting Models'."""
    par = doc.add_paragraph(style="Heading 2")
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(par, line=2.0, before=6, after=0)
    return par


def front_heading(doc, text, page_break=True):
    """Centred, bold front-matter heading (ABSTRACT, DEDICATION, ...)."""
    if page_break:
        doc.add_page_break()
    par = doc.add_paragraph(style="Heading 1")
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    _spacing(par, line=2.0, before=0, after=12)
    return par


def centered(doc, text, size=12, bold=False, italic=False, after=0, caps=False):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text.upper() if caps else text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return par


def blank(doc, n=1):
    for _ in range(n):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def equation(doc, text, number=None):
    """Centred equation; the number, if given, is flush with the right margin."""
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.italic = True
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(par, line=1.5, before=6, after=6)
    if number is not None:
        pf = par.paragraph_format
        pf.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
        tail = par.add_run(f"\t({number})")
        tail.font.name = FONT
        tail.font.size = BODY_PT
    return par


def definition(doc, term, meaning):
    """A bold term followed by its definition, used in Definition of Terms."""
    par = doc.add_paragraph()
    run = par.add_run(f"{term}: ")
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    run2 = par.add_run(meaning)
    run2.font.name = FONT
    run2.font.size = BODY_PT
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(par, line=2.0, after=0)
    return par


def abbreviation(doc, short, long):
    """Abbreviation and its expansion, aligned on a tab stop."""
    par = doc.add_paragraph()
    par.paragraph_format.tab_stops.add_tab_stop(Cm(4.0))
    run = par.add_run(f"{short}\t{long}")
    run.font.name = FONT
    run.font.size = BODY_PT
    _spacing(par, line=2.0, after=0)
    return par


def reference(doc, text):
    """APA 6 reference with a hanging indent."""
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-1.27)
    return par


# ── figures and tables ───────────────────────────────────────────────────────

def figure(doc, img_path, caption, width=Cm(15.5)):
    """Insert a figure with its caption BELOW, as the template requires."""
    path = Path(img_path)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(par, line=1.0, before=8, after=2)
    if path.exists():
        par.add_run().add_picture(str(path), width=width)
    else:
        par.add_run(f"[FIGURE NOT FOUND: {path.name}]")

    cap = doc.add_paragraph(style="FigureCaption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label, _, rest = caption.partition(":")
    r1 = cap.add_run(label + ":")
    r1.font.name = FONT
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = cap.add_run(rest)
    r2.font.name = FONT
    r2.font.size = Pt(11)
    _spacing(cap, line=1.0, before=2, after=12)
    return cap


def table_title(doc, caption):
    """Table title, placed ABOVE the table as the template requires."""
    par = doc.add_paragraph(style="TableCaption")
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label, _, rest = caption.partition(":")
    r1 = par.add_run(label + ":")
    r1.font.name = FONT
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = par.add_run(rest)
    r2.font.name = FONT
    r2.font.size = Pt(11)
    _spacing(par, line=1.0, before=10, after=2)
    return par


def table(doc, headers, rows, font_pt=10.5):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(text)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, line=1.0, after=0)
        r = p.runs[0]
        r.font.name = FONT
        r.font.size = Pt(font_pt)
        r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, line=1.0, after=0)
            r = p.runs[0]
            r.font.name = FONT
            r.font.size = Pt(font_pt)
    blank(doc)
    return tbl


# ── automatic lists (refresh with Ctrl+A then F9 in Word / WPS Writer) ───────

def toc_field(doc, instruction, hint):
    par = doc.add_paragraph()
    _spacing(par, line=1.5, after=6)
    _field(par, instruction, hint)
    return par


def table_of_contents(doc):
    toc_field(doc, r' TOC \o "1-2" \h \z \u ',
              "[Select the whole document and press F9 to build the table of "
              "contents.]")


def list_of_figures(doc):
    toc_field(doc, r' TOC \h \z \t "FigureCaption,1" ',
              "[Select the whole document and press F9 to build the list of "
              "figures.]")


def list_of_tables(doc):
    toc_field(doc, r' TOC \h \z \t "TableCaption,1" ',
              "[Select the whole document and press F9 to build the list of "
              "tables.]")


# ── appendix: source extraction and syntax-coloured rendering ───────────────

def extract(path, *names):
    """
    Pull named top-level functions/classes out of a source file, in the order
    requested. 'Class.method' selects a single method and prefixes it with the
    class header so the listing reads correctly.

    Reading the real files means the appendix cannot drift from the code that
    actually produced the results reported in Chapter Four.
    """
    src = Path(path).read_text(encoding="utf-8")
    lines = src.splitlines()
    tree = ast.parse(src)

    top = {n.name: n for n in tree.body
           if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef))}

    def block(node):
        start = min([node.lineno] + [d.lineno for d in node.decorator_list]) - 1
        return "\n".join(lines[start:node.end_lineno]).rstrip()

    chunks = []
    for name in names:
        if "." in name:
            cls_name, meth_name = name.split(".", 1)
            cls = top[cls_name]
            meth = next(n for n in cls.body
                        if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef))
                        and n.name == meth_name)
            chunks.append(f"class {cls_name}:\n" + block(meth))
        else:
            chunks.append(block(top[name]))
    return "\n\n\n".join(chunks)


def extract_methods(path, cls_name, names):
    """Several methods of one class under a single class header."""
    bodies = [extract(path, f"{cls_name}.{n}").split("\n", 1)[1] for n in names]
    return f"class {cls_name}:\n" + "\n\n".join(bodies)


def imports_of(path):
    """The import block of a module, for the head of an appendix listing."""
    src = Path(path).read_text(encoding="utf-8")
    lines = src.splitlines()
    tree = ast.parse(src)
    out = []
    for node in tree.body:
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            out.append("\n".join(lines[node.lineno - 1:node.end_lineno]))
    return "\n".join(out)


def _colour_spans(code):
    """Map each source line to a list of (start_col, end_col, colour) spans."""
    spans = {}

    def add(row, c0, c1, colour):
        spans.setdefault(row, []).append((c0, c1, colour))

    prev_kw = None
    try:
        tokens = list(tokenize.generate_tokens(io.StringIO(code).readline))
    except (tokenize.TokenError, IndentationError):
        return spans

    for tok in tokens:
        ttype, text, (r0, c0), (r1, c1), _ = tok
        if ttype == tokenize.COMMENT:
            colour = CLR_COMMENT
        elif ttype == tokenize.STRING:
            colour = CLR_STRING
        elif ttype == tokenize.NUMBER:
            colour = CLR_NUMBER
        elif ttype == tokenize.NAME:
            if keyword.iskeyword(text):
                colour = CLR_KEYWORD
            elif prev_kw == "def":
                colour = CLR_DEFNAME
            elif prev_kw == "class":
                colour = CLR_CLASSNAME
            elif text in BUILTIN_NAMES:
                colour = CLR_BUILTIN
            else:
                colour = None
        else:
            colour = None

        if ttype == tokenize.NAME:
            prev_kw = text if text in ("def", "class") else None
        elif ttype not in (tokenize.NL, tokenize.NEWLINE, tokenize.INDENT,
                           tokenize.DEDENT):
            prev_kw = None

        if colour is None:
            continue
        if r0 == r1:
            add(r0, c0, c1, colour)
        else:
            # A multi-line string: colour each line it covers.
            part = text.splitlines()
            for i, seg in enumerate(part):
                row = r0 + i
                start = c0 if i == 0 else 0
                add(row, start, start + len(seg), colour)
    return spans


def code_block(doc, code, first=True, last=True):
    """
    Render Python source as a shaded, syntax-coloured monospace block, one
    paragraph per line so it paginates cleanly.
    """
    code = code.replace("\t", "    ").rstrip("\n")
    spans = _colour_spans(code)
    lines = code.split("\n")

    for i, line in enumerate(lines, start=1):
        par = doc.add_paragraph()
        pf = par.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(6 if (i == 1 and first) else 0)
        pf.space_after = Pt(6 if (i == len(lines) and last) else 0)
        pf.left_indent = Cm(0.4)
        pf.right_indent = Cm(0.2)
        pf.keep_together = True
        _shade_paragraph(par, CODE_SHADE)

        pieces = []
        cursor = 0
        for c0, c1, colour in sorted(spans.get(i, [])):
            if c0 > cursor:
                pieces.append((line[cursor:c0], None))
            pieces.append((line[c0:c1], colour))
            cursor = max(cursor, c1)
        if cursor < len(line):
            pieces.append((line[cursor:], None))
        if not pieces:
            pieces = [(" ", None)]

        for text, colour in pieces:
            if text == "":
                continue
            run = par.add_run(text)
            run.font.name = CODE_FONT
            run.font.size = CODE_PT
            run.font.color.rgb = colour or CLR_PLAIN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    return doc


def appendix_heading(doc, label, caption, page_break=True):
    """'APPENDIX-A' followed by its one-line description (Busari layout)."""
    if page_break:
        doc.add_page_break()
    par = doc.add_paragraph(style="Heading 2")
    run = par.add_run(label)
    run.font.name = FONT
    run.font.size = Pt(13)
    run.bold = True
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(par, line=1.0, before=0, after=6)

    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.font.name = FONT
    run.font.size = BODY_PT
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(cap, line=1.5, before=0, after=8)
    return cap
